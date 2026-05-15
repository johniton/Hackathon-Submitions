"""
Hustlr Resume Builder — DOCX Service
Generates .docx resume from structured JSON using python-docx.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


def _add_heading(doc, text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def generate_docx(resume_json: dict) -> bytes:
    """Generate DOCX bytes from resume JSON."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume_json.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Headline
    if resume_json.get("headline"):
        hl = doc.add_paragraph()
        hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hl.add_run(resume_json["headline"])
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x64, 0x64, 0x78)

    # Contact line
    contact = []
    for k in ["email", "phone", "location", "linkedin", "github"]:
        if resume_json.get(k):
            contact.append(resume_json[k])
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(" | ".join(contact))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Summary
    if resume_json.get("summary"):
        _add_heading(doc, "Summary", level=2)
        p = doc.add_paragraph(resume_json["summary"])
        p.style.font.size = Pt(10)

    # Experience
    experience = resume_json.get("experience", [])
    if experience:
        _add_heading(doc, "Experience", level=2)
        for exp in experience:
            p = doc.add_paragraph()
            title_run = p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.bold = True
            title_run.font.size = Pt(10)
            if exp.get("dates"):
                p.add_run(f"  |  {exp['dates']}").font.size = Pt(9)

            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")
                for run in bp.runs:
                    run.font.size = Pt(9)

    # Education
    education = resume_json.get("education", [])
    if education:
        _add_heading(doc, "Education", level=2)
        for edu in education:
            p = doc.add_paragraph()
            run = p.add_run(edu.get("school", ""))
            run.bold = True
            run.font.size = Pt(10)
            if edu.get("dates"):
                p.add_run(f"  |  {edu['dates']}").font.size = Pt(9)
            if edu.get("degree"):
                dp = doc.add_paragraph(edu["degree"])
                dp.runs[0].font.size = Pt(9)

    # Projects
    projects = resume_json.get("projects", [])
    if projects:
        _add_heading(doc, "Projects", level=2)
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("name", ""))
            run.bold = True
            run.font.size = Pt(10)
            if proj.get("stack"):
                p.add_run(f"  ({', '.join(proj['stack'])})").font.size = Pt(9)
            if proj.get("description"):
                dp = doc.add_paragraph(proj["description"])
                dp.runs[0].font.size = Pt(9)

    # Skills
    skills = resume_json.get("skills", {})
    if skills:
        _add_heading(doc, "Skills", level=2)
        if isinstance(skills, dict):
            for category, items in skills.items():
                if isinstance(items, list) and items:
                    label = category.replace("_", " ").title()
                    p = doc.add_paragraph()
                    run = p.add_run(f"{label}: ")
                    run.bold = True
                    run.font.size = Pt(9)
                    p.add_run(", ".join(items)).font.size = Pt(9)
        elif isinstance(skills, list):
            p = doc.add_paragraph(", ".join(skills))
            p.runs[0].font.size = Pt(9)

    # Certifications
    certs = resume_json.get("certifications", [])
    if certs:
        _add_heading(doc, "Certifications", level=2)
        for c in certs:
            parts = [c.get("name", "")]
            if c.get("issuer"):
                parts.append(f"by {c['issuer']}")
            if c.get("date"):
                parts.append(f"({c['date']})")
            doc.add_paragraph(" — ".join(parts), style="List Bullet")

    # Achievements
    achievements = resume_json.get("achievements", [])
    if achievements:
        _add_heading(doc, "Achievements", level=2)
        for a in achievements:
            doc.add_paragraph(a, style="List Bullet")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
